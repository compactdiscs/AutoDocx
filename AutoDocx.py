import docx
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Mm
from docx.enum.text import *
import os
import sys

cn_num = ["一","二","三","四","五","六","七","八","九","十"]
num = ["1","2","3","4","5","6","7","8","9"]
punc = ["。","，","！","？","：","；","、",".","（","）"]
directory = os.path.dirname(os.path.realpath(sys.argv[0]))
w1 = directory + "/input.txt"
cnt = 0

#将txt按行读到list
f = open(w1,"r",encoding='utf-8')
#f = open("/Users/chendi/Documents/1工作/01模版/test/input.txt","r")
data = f.readlines()
data[-1] += "\n"
data.append("\n")
f.close()

length = len(data) #文件行数

#半角转全角
for k in range(0,length-1):
    data[k] = data[k].replace(",","，")
    data[k] = data[k].replace(";","；")
    data[k] = data[k].replace(":","：")
    data[k] = data[k].replace("!","！")
    data[k] = data[k].replace("?","？")
    data[k] = data[k].replace("(","（")
    data[k] = data[k].replace(")","）")
    data[k] = data[k].replace(" ","")
    data[k] = data[k].replace("\t", "")



#判断是否为附件或联系人格式
def isFJ(str):
    if (str[0:3] == "附件：") or (str[0:4] == "联系人："):
        return True
    else:
        return False

#判断是否为落款格式
'''
def isLK(str):
    if (str[0] in num) and (str[-2] == "日") and (len(str) <= 12):
        return True
    else:
        return False
'''
def isLK2(str):
    for i in str:
        if i in punc:
            return False
    return True

#判断是否为一级标题格式（如：一、xxx）
def is1BT(str):
    if (str[0] in cn_num) and (str[1] == "、"):
        return True
    else:
        return False

#判断是否为二级标题格式（如：（一）xxx）
def is2BT(str):
    if (str[0] == "（") and (str[1] in cn_num) and (str[2] == "）"):
        return True
    else:
        return False


w2 = directory + "/module.docx"
newfile = docx.Document(w2)

#newfile = docx.Document("/Users/chendi/Documents/1工作/01模版/test/module.docx")


for p in newfile.paragraphs:
    if p.text.lower() == '':  # 删除word中在开始部分的空白段落
        p = p._element
        p.getparent().remove(p)
        p._p = p._element = None


#处理头部空行
for j1 in range(0,length-1):
    if data[j1] != "\n":
        break


#标题（华文中宋、2号、加粗、居中、下端按2号字空一行）
p1 = newfile.add_paragraph()
p1.paragraph_format.line_spacing=Pt(30)  #行距固定值30磅
p1.paragraph_format.space_after = Pt(0)  #段后间距=0
text1 = p1.add_run("{} ".format(data[j1])) #下端按2号字空一行
text1.font.size = Pt(22)  # 字体大小2号
text1.bold = True  # 加粗
text1.font.name = '华文中宋'  # 控制是西文时的字体
text1.element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')  # 控制是中文时的字体
p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中


#正文若以"xxx单位："开头，需顶格
for j2 in range(j1+1,length-1):
    if (data[j2] != "\n") and (data[j2][-2] == "："):
        p2 = newfile.add_paragraph()
        p2.paragraph_format.line_spacing = Pt(30)  # 行距固定值30磅
        p2.paragraph_format.space_after = Pt(0)  # 段后间距=0
        text2 = p2.add_run("{}".format(data[j2][:-1]))
        text2.font.size = Pt(16)  # 字体大小3号
        text2.bold = False  # 字体不加粗
        text2.font.name = '仿宋'  # 控制是西文时的字体
        text2.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 控制是中文时的字体
        break
    elif (data[j2] == "\n"):
        continue
    else:
        j2 -= 1
        break

#处理正文
for i in range(j2+1,length-1):
    if data[i] == "\n": #忽略空白行
        continue
    elif isFJ(data[i]): #判断是否为附件或联系人格式
        str = "\n    " + data[i][:-1] #与正文空一行、首部空两格
    elif isLK2(data[i]): #判断是否为落款格式
        if cnt == 0:
            str = "\n" * 2 + " " * 38 + data[i][:-1]
            cnt+=1
        else:
            str = " " * 38 + data[i][:-1] #前置空格，顶到最右，需手动调整空格

    else: #普通正文格式
        str = "    " + data[i][:-1]

    p3 = newfile.add_paragraph()
    p3.paragraph_format.line_spacing = Pt(30)  # 行距固定值30磅
    p3.paragraph_format.space_after = Pt(0)  # 段后间距=0
    text3 = p3.add_run("{}".format(str))
    text3.font.size = Pt(16)  # 字体大小3号
    text3.bold = False  # 字体不加粗
    if is1BT(data[i]): #判断是否为一级标题格式（如：一、xxx）
        text3.font.name = '黑体'
        text3.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    elif is2BT(data[i]): #判断是否为二级标题格式（如：（一）xxx）
        text3.font.name = '楷体'
        text3.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    else: #普通正文格式
        text3.font.name = '仿宋'
        text3.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    '''
        elif isLK(data[i+1]): #判断是否为落款格式
            str = "\n" * 2 + " " * 38 + data[i][:-1] #前置空格，顶到最右，需手动调整空格
        elif isLK(data[i]): #判断是否为落款式
            str = " " * 38 + data[i][:-1] #前置空格，顶到最右，需手动调整空格
    '''

w3 = directory + "/{}.docx".format(data[j1][:-1])
newfile.save(w3)

#newfile.save("/Users/chendi/Desktop/{}.docx".format(data[j1][:-1]))

print("已生成《{}.docx》".format(data[j1][:-1]))
